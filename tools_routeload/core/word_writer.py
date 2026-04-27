"""把 FlightPlan 渲染为 Word 表格。

分组规则（v1.1）：
- 外层按航线 (起飞 ICAO, 目的 ICAO, 线路后缀) 分组 → 一个大标题
- 大标题下，按机型顺序 A319-115 → A320-214W → A320-251 依次放置该机型的载量分析表
"""
from __future__ import annotations

import json
from collections import defaultdict
from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .parser import FlightPlan

# 表头列定义：(显示名, FlightPlan 字段名)
COLUMNS = [
    ("月份", "month"),
    ("起飞重量\n（公斤）", "tow_kg"),
    ("总加油量\n（公斤）", "total_fuel_kg"),
    ("航程油量\n（公斤）", "trip_fuel_kg"),
    ("航程时间\n（时/分）", "trip_time"),
    ("航线距离\n（海里）", "trip_dist_nm"),
    ("最大业载\n（公斤）", "av_pld_kg"),
    ("航路平均风", "avg_wind"),
    ("额外油\n（公斤）", "extra_fuel_kg"),
    ("落地剩油\n（公斤）", "target_arrival_kg"),
    ("计算高度\n(FT)", "calc_alt_ft"),
    ("限重计算温度\n(℃)", "_zero"),
    ("人数", "pax_count"),
]

# 机型在同一航线下的展示顺序
AIRCRAFT_ORDER = ["A319-115", "A320-214", "A320-214W", "A320-251", "A321-211"]

CONFIG_DIR = Path(__file__).resolve().parent.parent / "config"


def _load_json(name: str) -> dict:
    p = CONFIG_DIR / name
    if not p.exists():
        return {}
    data = json.loads(p.read_text(encoding="utf-8"))
    return {k: v for k, v in data.items() if not k.startswith("_")}


def _airport_name(icao: str, mapping: dict) -> str:
    return mapping.get(icao, icao)


def _aircraft_name(code: str, mapping: dict) -> str:
    return mapping.get(code, code or "未知机型")


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = p.add_run(str(text))
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def _set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "000000")
        tc_borders.append(b)


def _set_table_layout_fixed(table):
    """固定表格布局，列宽不会因为单元格内容自动扩张。"""
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def _value_for(fp: FlightPlan, attr: str):
    if attr == "_zero":
        return 0
    if attr == "calc_alt_ft":
        return fp.calc_alt_ft
    return getattr(fp, attr)


def _aircraft_sort_key(code: str, aircraft_map: dict) -> tuple[int, str]:
    name = _aircraft_name(code, aircraft_map)
    if name in AIRCRAFT_ORDER:
        return (AIRCRAFT_ORDER.index(name), name)
    return (len(AIRCRAFT_ORDER), name)


def _render_table(doc: Document, items: list[FlightPlan], ac_code: str, aircraft_map: dict, page_width_cm: float):
    """渲染单个机型的载量分析表。"""
    items = sorted(items, key=lambda x: x.month)

    # 副标题：湖南航空公司A319-115飞机航线及载量分析
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(2)
    sub.paragraph_format.space_after = Pt(2)
    run = sub.add_run(f"湖南航空公司{_aircraft_name(ac_code, aircraft_map)}飞机航线及载量分析")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    n_cols = len(COLUMNS)
    table = doc.add_table(rows=2 + len(items), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_layout_fixed(table)

    # 列宽：纵向 A4 内宽 ≈ 18cm，13 列均分 ≈ 1.38cm
    col_w = page_width_cm / n_cols
    for col in table.columns:
        for cell in col.cells:
            cell.width = Cm(col_w)

    # 第 1 行：信息合并
    info_row = table.rows[0]
    info_row.cells[0].merge(info_row.cells[-1])
    sample = items[0]
    info = (
        f"机号：{sample.aircraft_reg}; "
        f"起飞机场：{sample.dep_icao}; "
        f"目的地机场：{sample.arr_icao}; "
        f"备降场：{sample.altn_icao}; "
        f"备降距离：{sample.altn_dist_nm}NM"
    )
    _set_cell_text(info_row.cells[0], info, bold=True, size=9)

    # 第 2 行：表头
    header_row = table.rows[1]
    for i, (name, _) in enumerate(COLUMNS):
        _set_cell_text(header_row.cells[i], name, bold=True, size=8)

    # 数据行
    for r, fp in enumerate(items, start=2):
        row = table.rows[r]
        for i, (_, attr) in enumerate(COLUMNS):
            _set_cell_text(row.cells[i], _value_for(fp, attr), size=9)

    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell)


def build_doc(
    plans: Iterable[FlightPlan],
    airports: dict | None = None,
    aircraft: dict | None = None,
) -> Document:
    """渲染。外层航线分组（含线路后缀），内层按机型顺序。"""
    if airports is None:
        airports = _load_json("airports.json")
    if aircraft is None:
        aircraft = _load_json("aircraft.json")

    # 外层：航线
    route_groups: dict[tuple, dict[str, list[FlightPlan]]] = defaultdict(lambda: defaultdict(list))
    route_order: list[tuple] = []
    for fp in plans:
        rkey = (fp.dep_icao, fp.arr_icao, fp.route_suffix)
        if rkey not in route_groups:
            route_order.append(rkey)
        route_groups[rkey][fp.aircraft_type_code].append(fp)

    # 往返配对排序：每条航线之后紧邻其反向航线（同线路后缀），按首次出现顺序为基准。
    paired_order: list[tuple] = []
    placed: set[tuple] = set()
    for rkey in route_order:
        if rkey in placed:
            continue
        paired_order.append(rkey)
        placed.add(rkey)
        dep, arr, suffix = rkey
        rev = (arr, dep, suffix)
        if rev in route_groups and rev not in placed:
            paired_order.append(rev)
            placed.add(rev)
    route_order = paired_order

    doc = Document()
    section = doc.sections[0]
    # 纵向 A4：21cm × 29.7cm（python-docx 默认就是 portrait）
    from docx.shared import Cm as _Cm
    section.page_width = _Cm(21.0)
    section.page_height = _Cm(29.7)
    section.left_margin = section.right_margin = _Cm(1.5)
    section.top_margin = section.bottom_margin = _Cm(1.5)
    page_inner_w = 21.0 - 1.5 * 2  # 18cm

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    for idx, rkey in enumerate(route_order, 1):
        dep, arr, suffix = rkey
        ac_groups = route_groups[rkey]

        # 大标题：1. 无锡-吐鲁番（南线）
        title_text = f"{idx}. {_airport_name(dep, airports)}-{_airport_name(arr, airports)}"
        if suffix:
            title_text += f"（{suffix}）"
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_p.paragraph_format.space_before = Pt(6)
        title_p.paragraph_format.space_after = Pt(2)
        run = title_p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 机型按预设顺序
        ac_codes = sorted(ac_groups.keys(), key=lambda c: _aircraft_sort_key(c, aircraft))
        for ac_code in ac_codes:
            _render_table(doc, ac_groups[ac_code], ac_code, aircraft, page_inner_w)

        # 段间距
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

    return doc


def build_bytes(plans: Iterable[FlightPlan], **kw) -> bytes:
    doc = build_doc(plans, **kw)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
